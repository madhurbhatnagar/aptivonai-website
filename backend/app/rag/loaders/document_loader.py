from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader


def extract_text(file_path: Path, extension: str) -> str:
    extension = extension.lower().lstrip(".")

    if extension == "pdf":
        return _extract_pdf(file_path)
    if extension == "docx":
        return _extract_docx(file_path)
    if extension == "txt":
        return _extract_txt(file_path)
    if extension == "csv":
        return _extract_csv(file_path)
    if extension == "xlsx":
        return _extract_xlsx(file_path)

    raise ValueError(f"Unsupported file extension: {extension}")


def _extract_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = []

    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {page_number}]\n{text.strip()}")

    return "\n\n".join(pages)


def _extract_docx(file_path: Path) -> str:
    document = Document(str(file_path))
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append(f"[Table {table_index}]\n" + "\n".join(rows))

    return "\n\n".join(parts)


def _extract_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def _extract_csv(file_path: Path) -> str:
    dataframe = pd.read_csv(file_path, dtype=str).fillna("")
    return dataframe.to_csv(index=False)


def _extract_xlsx(file_path: Path) -> str:
    sheets = pd.read_excel(file_path, sheet_name=None, dtype=str)
    parts = []

    for sheet_name, dataframe in sheets.items():
        clean_dataframe = dataframe.fillna("")
        parts.append(f"[Sheet: {sheet_name}]\n{clean_dataframe.to_csv(index=False)}")

    return "\n\n".join(parts)
